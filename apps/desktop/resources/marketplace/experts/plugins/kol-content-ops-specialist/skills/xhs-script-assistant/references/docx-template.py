#!/usr/bin/env python3
"""
小红书脚本审核 Word 报告模板生成器（xhs-script-assistant 通用）
用法：python3 docx-template.py <json_file> [output_path]
输入：JSON 文件包含审核报告全部内容
输出：带样式的 .docx 文件
"""

import json, sys, os, re
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("缺少 python-docx，无法生成 Word；请使用桌面应用打包运行时，或改用 Markdown 输出模式。", file=sys.stderr)
    raise SystemExit(2)

# ===== Color constants =====
RED = RGBColor(0xCF, 0x13, 0x22)
ORANGE = RGBColor(0xD4, 0x6B, 0x08)
GREEN = RGBColor(0x38, 0x9E, 0x0D)

# ===== 文件命名：内容阶段-品牌产品名-博主名称 =====
def safe_name(s):
    """去除文件名非法字符（/ \\ : * ? " < > |），并去首尾空格。"""
    return re.sub(r'[/\\:*?"<>|]', '', str(s)).strip()

def make_docx_name(stage, product, talent):
    """生成统一 Word 文件名：【内容阶段-品牌产品名-博主名称】.docx"""
    st = safe_name(stage) or "视频脚本"
    pd = safe_name(product) or "未命名产品"
    tl = safe_name(talent) or "达人"
    return f"【{st}-{pd}-{tl}】.docx"
MUTED = RGBColor(0x94, 0xA3, 0xB8)
ACCENT = RGBColor(0xD4, 0x38, 0x0D)

LEVEL_COLORS = {"P0": RED, "P1": ORANGE, "P2": GREEN}

# ===== Helpers =====
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_run(para, text, bold=False, color=None, size=None, strike=False):
    run = para.add_run(text)
    run.bold = bold
    if color: run.font.color.rgb = color
    if size: run.font.size = size
    if strike: run.font.strike = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return run

def render_cell(cell, data):
    """Render cell data: string or list of {text, bold, color, strike} dicts"""
    p = cell.paragraphs[0]
    if isinstance(data, str):
        add_run(p, data, size=Pt(10))
    elif isinstance(data, list):
        for seg in data:
            color = None
            if seg.get("color"):
                c = seg["color"]
                if c == "RED": color = RED
                elif c == "ORANGE": color = ORANGE
                elif c == "GREEN": color = GREEN
                elif c == "MUTED": color = MUTED
                elif c == "ACCENT": color = ACCENT
            add_run(p, seg["text"],
                    bold=seg.get("bold", False),
                    color=color,
                    size=Pt(10),
                    strike=seg.get("strike", False))

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, 'F1F5F9')
        p = cell.paragraphs[0]
        add_run(p, h, bold=True, size=Pt(10))
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            render_cell(table.rows[r_idx+1].cells[c_idx], cell_data)
    # Widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def render_risk_level(level_str):
    """Convert 'P0'/'P1'/'P2' to styled run dict"""
    return [{"text": level_str, "bold": True, "color": level_str}]

# ===== Main builder =====
def build_report(data, output_path):
    doc = Document()

    # Style
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Section titles (allow override for batch/slim mode)
    # Pass data["titles"] = {"script": "一、修改后脚本", "brief_check": "二、审核核对", ...}
    T = data.get("titles", {})
    def title_of(key, default):
        return T.get(key, default)

    # Title
    title = doc.add_heading(data.get("title", "脚本审核报告"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = data.get("meta", "")
    add_run(meta, meta_text, size=Pt(10))

    doc.add_paragraph()

    # ---- Brief提取 ----
    if "brief_rows" in data:
        doc.add_heading(title_of("brief", "Brief 项目要求提取"), level=1)
        add_table(doc, ["字段", "内容"], data["brief_rows"], col_widths=[3, 14])
        doc.add_paragraph()

    # ---- 区块一：修改后脚本 ----
    if "script_rows" in data:
        doc.add_heading(title_of("script", "一、修改后脚本（含留痕）"), level=1)

        if "p0_warning" in data:
            warn = doc.add_paragraph()
            add_run(warn, "⚠ P0 ", bold=True, color=RED, size=Pt(12))
            add_run(warn, data["p0_warning"], size=Pt(11))
            doc.add_paragraph()

        add_table(doc,
                  ["序号", "时间", "画面内容", "口播/字幕", "花字"],
                  data["script_rows"],
                  col_widths=[1.2, 1.5, 3.5, 9, 2])
        doc.add_paragraph()

    # 发布文案
    if "pub_copy" in data:
        p = doc.add_paragraph()
        pc = data["pub_copy"]
        add_run(p, "📌 发布文案建议修改：\n", bold=True, size=Pt(10))
        add_run(p, "原：" + pc["original"] + "\n", size=Pt(10))
        add_run(p, "改：", bold=True, size=Pt(10))
        add_run(p, pc["modified"], bold=True, color=RED, size=Pt(10))
        add_run(p, "\n理由：" + pc["reason"], size=Pt(10))

    # 话题
    if "topic" in data:
        t = doc.add_paragraph()
        tp = data["topic"]
        add_run(t, "📌 话题建议修改：\n", bold=True, size=Pt(10))
        add_run(t, "原：" + tp["original"] + "\n", size=Pt(10))
        add_run(t, "改：", bold=True, size=Pt(10))
        add_run(t, tp["modified"], bold=True, color=RED, size=Pt(10))
        add_run(t, "\n理由：" + tp["reason"], size=Pt(10))

    # ---- 区块二：修改说明 ----
    if "mod_rows" in data:
        doc.add_heading(title_of("mod", "二、修改说明"), level=1)
        # Convert risk level strings to styled dicts
        for row in data["mod_rows"]:
            if len(row) > 4 and isinstance(row[4], str) and row[4] in LEVEL_COLORS:
                row[4] = render_risk_level(row[4])
        add_table(doc,
                  ["位置", "修改类型", "修改原因", "对应依据", "风险等级"],
                  data["mod_rows"],
                  col_widths=[2.5, 1.8, 4.5, 3.2, 1.5])
        doc.add_paragraph()

    # ---- 区块三：Brief符合度 ----
    if "brief_check_rows" in data:
        doc.add_heading(title_of("brief_check", "三、Brief 符合度检查"), level=1)
        add_table(doc,
                  ["Brief 要求", "修改后是否满足", "对应脚本位置", "备注"],
                  data["brief_check_rows"],
                  col_widths=[4, 2, 2.5, 5])
        doc.add_paragraph()

    # ---- 区块四：风险表达（legacy，已被 issue_rows 合并）----
    if "risk_rows" in data:
        doc.add_heading(title_of("risk", "四、风险表达处理"), level=1)
        add_table(doc,
                  ["原表达", "风险原因", "修改后表达"],
                  data["risk_rows"],
                  col_widths=[4.5, 4, 4.5])
        doc.add_paragraph()

    # ---- 审核核对：统一问题表（6 字段，canonical）----
    if "issue_rows" in data:
        doc.add_heading(title_of("issues", "审核核对（问题表）"), level=1)
        # 风险等级自动着色：P0/P1/P2
        for row in data["issue_rows"]:
            if len(row) > 3 and isinstance(row[3], str) and row[3] in LEVEL_COLORS:
                row[3] = render_risk_level(row[3])
        add_table(doc,
                  ["问题位置", "问题类型", "具体问题", "风险等级", "修改建议", "可替代表达"],
                  data["issue_rows"],
                  col_widths=[2, 2, 4, 1.5, 4, 4])
        doc.add_paragraph()

    # ---- 区块五：参考案例 ----
    if "ref_case" in data:
        doc.add_heading(title_of("ref_case", "五、参考案例调用"), level=1)
        rc = data["ref_case"]
        for para_data in rc.get("paragraphs", []):
            p = doc.add_paragraph()
            for seg in para_data:
                color = None
                if seg.get("color"):
                    c = seg["color"]
                    if c == "RED": color = RED
                    elif c == "ORANGE": color = ORANGE
                    elif c == "GREEN": color = GREEN
                add_run(p, seg["text"], bold=seg.get("bold", False), color=color, size=Pt(10))
        doc.add_paragraph()

    # ---- 区块六：遗漏自查 ----
    if "omit_rows" in data:
        doc.add_heading(title_of("omit", "六、遗漏自查"), level=1)
        add_table(doc,
                  ["缺失项", "影响", "建议"],
                  data["omit_rows"],
                  col_widths=[3.5, 6, 4.5])
        doc.add_paragraph()

    # ---- 质检清单 ----
    if "checks" in data:
        doc.add_heading(title_of("checks", "质检自查清单"), level=1)
        for check in data["checks"]:
            p = doc.add_paragraph()
            icon = check.get("icon", "✅")
            color = ORANGE if icon == "⚠️" else GREEN
            add_run(p, f"{icon} ", bold=True, color=color, size=Pt(10))
            add_run(p, check["text"], size=Pt(10))

    doc.save(output_path)
    print(f"Saved: {output_path}")


def build_batch(data_list, output_dir, project_name="", stage=""):
    """批量生成：每个达人一个 Word。文件名格式：【内容阶段-品牌产品名-博主名称】.docx"""
    saved = []
    used_paths = set()
    for data in data_list:
        talent = data.get("talent", "达人")
        product = data.get("product", project_name) or project_name
        st = data.get("stage", stage) or stage or "视频脚本"
        fname = make_docx_name(st, product, talent)
        out_path = os.path.join(output_dir, fname)
        stem, extension = os.path.splitext(out_path)
        suffix = 2
        while out_path in used_paths or os.path.exists(out_path):
            out_path = f"{stem}_{suffix}{extension}"
            suffix += 1
        build_report(data, out_path)
        saved.append(out_path)
        used_paths.add(out_path)
    return saved

# ===== Entry =====
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage:")
        print("  单篇: python3 docx-template.py <json_file> [output_path]")
        print("  批量: python3 docx-template.py --batch <json_file> <output_dir>")
        sys.exit(1)

    if sys.argv[1] == "--batch":
        # 批量模式：JSON 为 {"project": "项目名", "reports": [data1, data2, ...]}
        json_path = sys.argv[2]
        output_dir = sys.argv[3] if len(sys.argv) > 3 else os.path.dirname(json_path)
        with open(json_path, "r", encoding="utf-8") as f:
            batch = json.load(f)
        reports = batch.get("reports", [])
        project = batch.get("project", "")
        stage = batch.get("stage", "")
        saved = build_batch(reports, output_dir, project, stage)
        print(f"\n批量完成，共生成 {len(saved)} 份：")
        for s in saved:
            print(f"  - {s}")
    else:
        json_path = sys.argv[1]
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if len(sys.argv) > 2:
            output_path = sys.argv[2]
        else:
            talent = data.get("talent", "达人")
            product = data.get("product", data.get("project", ""))
            stage = data.get("stage", "视频脚本")
            output_path = os.path.join(os.path.dirname(json_path), make_docx_name(stage, product, talent))
        build_report(data, output_path)
