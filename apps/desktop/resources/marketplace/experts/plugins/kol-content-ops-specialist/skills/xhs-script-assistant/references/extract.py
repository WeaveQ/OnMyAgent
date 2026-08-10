#!/usr/bin/env python3
"""
通用脚本/Brief 提取器 —— 支持 word/excel/pdf 自动识别
用途：代替模型逐个阅读文件，脚本化提取干净文本，大幅降低积分消耗
用法：
    python3 extract.py <file_or_dir> [--out <output.json>]
    - 传单个文件：提取该文件
    - 传目录：批量提取目录下所有 .xlsx/.docx/.pdf
输出：结构化 JSON（文件名 → 提取内容），或直接打印
"""

import sys, os, json, glob

def extract_excel(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    result = {}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None]
            if vals:
                rows.append(vals)
        result[sheet_name] = rows
    return {"type": "excel", "sheets": result}

def extract_docx(path):
    from docx import Document
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for tbl in doc.tables:
        t_rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                t_rows.append(cells)
        if t_rows:
            tables.append(t_rows)
    return {"type": "docx", "paragraphs": paras, "tables": tables}

def extract_pdf(path):
    import PyPDF2
    reader = PyPDF2.PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
            text = text.encode('utf-8', errors='replace').decode('utf-8')
            pages.append({"page": i + 1, "text": text})
        except Exception as e:
            pages.append({"page": i + 1, "text": f"[EXTRACT ERROR: {e}]"})
    return {"type": "pdf", "pages": pages}

def extract_file(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in (".xlsx", ".xlsm"):
            return extract_excel(path)
        elif ext == ".docx":
            return extract_docx(path)
        elif ext == ".pdf":
            return extract_pdf(path)
        else:
            return {"type": "unsupported", "note": f"跳过不支持的格式：{ext}"}
    except ImportError as error:
        dependency = getattr(error, "name", None) or "文档解析依赖"
        return {
            "type": "dependency_error",
            "note": f"缺少 {dependency}，请使用桌面应用打包运行时；不会自动安装或修改系统 Python。",
        }

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 extract.py <file_or_dir> [--out <output.json>]")
        sys.exit(1)

    target = sys.argv[1]
    out_path = None
    if "--out" in sys.argv:
        out_path = sys.argv[sys.argv.index("--out") + 1]

    result = {}
    if os.path.isdir(target):
        files = []
        for ext in ("*.xlsx", "*.xlsm", "*.docx", "*.pdf"):
            files.extend(glob.glob(os.path.join(target, ext)))
        for f in sorted(files):
            if os.path.basename(f).startswith("~$"):  # skip temp files
                continue
            result[os.path.basename(f)] = extract_file(f)
    else:
        result[os.path.basename(target)] = extract_file(target)

    output = json.dumps(result, ensure_ascii=False, indent=2)
    if out_path:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(output)
        print(f"Extracted {len(result)} file(s) → {out_path}")
    else:
        print(output)

if __name__ == "__main__":
    main()
